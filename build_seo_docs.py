"""
Build two audience-tailored Word docs from the INTERNATIONAL-SEO-PLAN.md.

- Apurvi-SEO-Team-Brief.docx       (strategy / keywords / content / links / paid / calendar)
- Apurvi-Tech-Team-Brief.docx       (on-site code, hreflang, schema, AEO format, perf)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ------------------------------- styling helpers -------------------------------

NAVY = RGBColor(0x3F, 0x3A, 0x8A)        # Apurvi brand colour
GREY_DARK = RGBColor(0x33, 0x33, 0x33)
GREY_MID = RGBColor(0x6D, 0x6E, 0x71)
GREY_BG = "F2F2F2"
ACCENT = RGBColor(0xC0, 0x39, 0x2B)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def setup_styles(doc):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = GREY_DARK

    for h, size in (("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13)):
        st = doc.styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY


def add_cover(doc, title, subtitle, audience):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("APURVI INDUSTRIES")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.font.size = Pt(14)
    r.font.color.rgb = GREY_MID

    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run(f"Prepared for: {audience}")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = GREY_DARK

    p = doc.add_paragraph()
    r = p.add_run("Date: 29 June 2026  ·  Target geos: United States + United Kingdom")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY_MID

    p = doc.add_paragraph()
    r = p.add_run("Compiled from 6 parallel deep-research agents covering keywords, content, authority/links, technical/AI search, competitor teardown, and impressions math.")
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GREY_MID
    doc.add_paragraph("")


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_h3(doc, text):
    return doc.add_heading(text, level=3)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            r = p.add_run(item[0])
            r.bold = True
            p.add_run(" — " + item[1])
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, label, body):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, GREY_BG)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(body)
    doc.add_paragraph("")


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], "3F3A8A")
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph("")


def add_code_block(doc, code):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F4F4F4")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = GREY_DARK
    doc.add_paragraph("")


def add_page_break(doc):
    doc.add_page_break()


# =============================================================================
#                     DOC A — SEO TEAM BRIEF
# =============================================================================

def build_seo_doc():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    add_cover(
        doc,
        title="US + UK SEO Master Plan",
        subtitle="400,000 Search Impressions in 30 Days  —  SEO & Marketing Team Brief",
        audience="SEO Team & Marketing",
    )
    add_page_break(doc)

    # 1. EXECUTIVE SUMMARY
    add_h1(doc, "1. Executive Summary")
    add_para(
        doc,
        "Apurvi Industries is launching a 30-day US + UK SEO push targeting 400,000 search impressions. "
        "This brief is the SEO and marketing team's complete playbook: keywords to target, content to publish, "
        "links to acquire, paid campaigns to launch, and a day-by-day calendar. The tech team has a parallel "
        "brief covering on-site code changes — work in lock-step with them.",
    )
    add_callout(
        doc,
        "Brutal-honest verdict",
        "400K from pure organic in 30 days is impossible for a new US/UK entrant. 400K hybrid (organic + AI search + "
        "community + ~$8–10K paid) is achievable. The smarter target is 250–400K impressions and 30 qualified "
        "RFQs in month 1, then scale to 600K–1M in months 2–3 as organic compounds.",
    )

    add_h2(doc, "Apurvi's unfair advantages in 2026")
    add_bullets(doc, [
        ("Section 232 tariffs", "Trump's April 6, 2026 proclamation raised Section 232 to 50% on Chinese stainless. US procurement teams are urgently searching for non-Chinese suppliers right now."),
        ("UK CBAM phase-in", "Forces European buyers to verify origin and carbon credentials. EN 10204 3.1 MTC — which Apurvi already issues — is now mandatory."),
        ("Aluminum vapour recovery tubing", "Fewer than 5 manufacturers globally serve this niche. Near-monopoly opportunity."),
        ("Real Mehsana plant + ISO 9001:2015 + 25–200 HP coverage", "Unbeatable E-E-A-T versus paper-thin competitors."),
        ("AI search timing", "Google AI Overviews and Bing Copilot now appear in 35–40% of US/UK B2B queries. No Indian competitor is doing AEO well — Apurvi can own this for 6–12 months."),
    ])

    add_h2(doc, "Impression-build formula (30-day total)")
    add_table(
        doc,
        ["Layer", "30-day impressions", "Effort", "Cost"],
        [
            ["Organic (long-tail + brand)", "65–130K", "20 new pages + on-site fixes", "Time only"],
            ["AI search citations (AIO / Perplexity / ChatGPT / Bing Copilot)", "15–30K", "AEO content + llms.txt", "Time only"],
            ["LinkedIn organic (founder + company)", "30–100K", "5 posts/week + 1 newsletter", "Time only"],
            ["Reddit + Quora organic", "20–80K", "3–5 posts/week", "Time only"],
            ["YouTube organic (factory shorts)", "5–20K", "4–6 videos", "Phone + 4 hrs/video"],
            ["Pinterest + niche directories", "5–15K", "30 pins + profile claims", "8 hrs total"],
            ["Paid (Google + Bing + LinkedIn + Reddit + YouTube + Display + Quora)", "200–300K", "Campaign setup + mgmt", "~$8–10K"],
            ["TOTAL", "~420K (range 340–675K)", "", "$8–10K"],
        ],
    )

    add_page_break(doc)

    # 2. KEYWORDS
    add_h1(doc, "2. Top 30 US + UK Keywords to Target")
    add_para(doc, "Volumes are monthly. KD = Keyword Difficulty (0–100). Source: blended estimates from Google Keyword Planner banding, Ahrefs/SEMrush public data, expert reasoning.")
    add_table(
        doc,
        ["#", "Keyword", "US Vol", "UK Vol", "KD", "Target page"],
        [
            ["1",  "stainless steel pipe suppliers", "2,400", "880", "58", "/us/ or /uk/ homepage"],
            ["2",  "ASTM A312 pipe supplier", "590", "110", "32", "/us/astm-a312-tp304-stainless-steel-pipe-supplier"],
            ["3",  "ASTM A312 TP304 seamless pipe", "480", "90", "28", "same as above"],
            ["4",  "ASTM A268 stainless tube", "320", "70", "18", "/us/astm-a268-stainless-steel-tubing-manufacturer"],
            ["5",  "stainless steel motor shell tube", "90", "40", "12", "/us/submersible-pump-motor-shell-tube-supplier (HERO)"],
            ["6",  "submersible pump motor housing tube", "170", "70", "14", "same as above"],
            ["7",  "EN 10296 welded stainless tube", "210", "320", "24", "/uk/en-10296-2-welded-stainless-steel-tube-supplier"],
            ["8",  "EN 10204 3.1 MTC supplier", "140", "260", "22", "/uk/en-10204-3-1-mtc-stainless-pipe-manufacturer"],
            ["9",  "stainless steel pipe wholesaler USA", "320", "—", "36", "/us/"],
            ["10", "stainless steel tube manufacturer India", "880", "390", "41", "/company → /us/about-us-and-export"],
            ["11", "SS 304 vs 316", "6,600", "2,400", "26", "Extend existing blog with US/UK angle"],
            ["12", "ASTM A312 vs A213", "590", "170", "14", "New blog (quick-win)"],
            ["13", "aluminum vapor recovery tubing", "90", "—", "8", "/us/aluminum-vapor-recovery-tubing-stage-i-retrofit (HERO)"],
            ["14", "aluminium vapour recovery pipe", "—", "70", "6", "/uk/aluminium-vapour-recovery-tubing"],
            ["15", "EPA Stage II vapor recovery piping", "260", "—", "22", "/us/aluminum-vapor-recovery-tubing-stage-i-retrofit"],
            ["16", "6061 aluminum tubing supplier", "720", "140", "38", "/us/aluminum-vapor-recovery-tubing"],
            ["17", "6063 aluminium tube UK", "50", "210", "18", "/uk/aluminium-vapour-recovery-tubing"],
            ["18", "precision aluminium tubing UK", "90", "480", "28", "/uk/precision-aluminium-tubing"],
            ["19", "borehole pump tube supplier", "30", "170", "12", "/uk/borehole-pump-tube-supplier"],
            ["20", "submersible pump pipe manufacturer", "480", "90", "30", "/us/submersible-pump-motor-shell-tube-supplier"],
            ["21", "NRV pipe stainless steel", "110", "90", "14", "/products#pump-pipe-nrv"],
            ["22", "25 HP submersible pump tube", "40", "20", "6", "/us/high-hp-submersible-motor-tube"],
            ["23", "100 HP submersible motor tube manufacturer", "30", "10", "4", "same as above"],
            ["24", "stainless steel tube for OEM", "170", "90", "22", "/us/stainless-steel-tube-for-oem"],
            ["25", "seamless stainless steel pipe USA", "590", "—", "44", "/us/"],
            ["26", "ISO 9001 stainless tube supplier", "110", "70", "16", "/quality-certifications"],
            ["27", "ground OD ID stainless tube", "50", "30", "8", "/products#pressure-pump-sleeve"],
            ["28", "pump shaft sleeve tubing 316L", "70", "30", "10", "/products#pressure-pump-sleeve"],
            ["29", "long stainless steel tube 6m", "40", "30", "6", "/products#manufacturing-long-tube"],
            ["30", "deep well submersible motor tube", "90", "40", "8", "/us/high-hp-submersible-motor-tube"],
        ],
    )

    add_h2(doc, "10 quick-win keywords (rank top-10 in < 60 days)")
    add_para(doc, "Long-tail B2B queries no Western competitor owns. Each gets a 700–1,200 word landing page with the exact keyword in H1, slug, and first 100 words.")
    add_numbered(doc, [
        "stainless steel motor tube manufacturer for OEM submersible pumps (KD 3)",
        "EN 10204 3.1 MTC stainless steel motor tube (KD 4)",
        "aluminum vapour recovery tubing manufacturer India (KD 2)",
        "ASTM A268 TP409 motor shell tube supplier (KD 3)",
        "25 HP to 200 HP submersible pump tube (KD 3)",
        "ground OD ID stainless tube for pump sleeve (KD 2)",
        "6063-T6 aluminium vapour tubing UK (KD 4)",
        "ASTM A312 TP316L motor tube wholesale price (KD 6)",
        "submersible pump pipe with NRV thread (KD 3)",
        "precision stainless steel tube 2 inch to 12 inch OD (KD 5)",
    ])

    add_h2(doc, "US vs UK spelling rules — CRITICAL")
    add_callout(doc, "Why this matters", "Apurvi currently writes 'Aluminum Vapour Tubing' — a hybrid that ranks for neither market. Every /us/ page must use US spelling; every /uk/ page must use UK spelling.")
    add_table(
        doc,
        ["Concept", "US spelling", "UK spelling"],
        [
            ["Metal", "aluminum", "aluminium"],
            ["Vapor", "vapor", "vapour"],
            ["Borewell / water well / borehole", "water well", "borehole"],
            ["Pump pipe / pump tube", "pump pipe", "pump tube / tubing"],
            ["Catalog / catalogue", "catalog", "catalogue"],
        ],
    )

    add_page_break(doc)

    # 3. CONTENT
    add_h1(doc, "3. Content Architecture — 3 Pillars × 8 Clusters = 27 New Pages")
    add_para(doc, "Each pillar is ~3,500 words. Each cluster is ~1,500–2,500 words. Every cluster links UP to its pillar and SIDEWAYS to 2 sibling clusters.")

    add_h2(doc, "PILLAR A — Stainless Steel Motor Tubes for Submersible Pumps (US/UK Buyer's Guide)")
    add_bullets(doc, [
        ("A1", "304L vs 316L for submersible motor housings (saltwater/freshwater decision tree)"),
        ("A2", "4″ vs 6″ vs 8″ motor tube sizing — HP / GPM / depth mapping"),
        ("A3", "ASTM A312 vs A213 — which standard your pump motor tube actually needs"),
        ("A4", "Wall thickness selection: pressure rating + collapse depth calc"),
        ("A5", "Welded vs seamless motor tubes — UL/NSF pump assembler acceptance criteria"),
        ("A6", "Pickling, passivation & 0.4 Ra ID finish — why ID finish kills pump life"),
        ("A7", "Higher HP (25–200 HP) deep-well motor tube — long-tube manufacturing explained"),
        ("A8", "Avoiding pitting corrosion in coastal US/UK groundwater — chloride field data"),
    ])

    add_h2(doc, "PILLAR B — Aluminum Vapour Recovery Tubing for Gas Stations (EPA + UK PPC)")
    add_bullets(doc, [
        ("B1", "Al 6061 vs 6063 vs 5052 — alloy for vapour recovery lines"),
        ("B2", "Stage II decommissioning in the US — what tubing you keep vs replace"),
        ("B3", "UK PPC permits + vapour recovery — new-build material spec"),
        ("B4", "EN 754 vs ASTM B210 for aluminum drawn tubing — cross-reference"),
        ("B5", "Forecourt fire-safety + aluminum tubing weld practices"),
        ("B6", "Replacing copper vapour lines with aluminum — TCO calculator"),
        ("B7", "Sourcing aluminum vapour tubing from India — HS codes 7608, lead times"),
        ("B8", "UL/ETV listing pathway for imported aluminum tubing"),
    ])

    add_h2(doc, "PILLAR C — Sourcing SS Pipe from India: Risk-Free Procurement Playbook")
    add_bullets(doc, [
        ("C1", "EN 10204 Type 3.1 MTC — how to read it (annotated sample)"),
        ("C2", "ISO 9001:2015 + IS:6761 audit checklist before signing an Indian supplier"),
        ("C3", "SS 202 vs 304 vs 316 — buyer downgrade traps"),
        ("C4", "Incoterms: CIF vs FOB to US East Coast / UK Felixstowe"),
        ("C5", "US Section 232 + UK CBAM on Indian SS — 2026 landed-cost reality"),
        ("C6", "Third-party inspection (SGS/TUV/BV) — when it pays back"),
        ("C7", "SS pressure pump sleeves + NRV pipe — buyer spec template"),
        ("C8", "ASTM A312 mill traceability — heat number → MTC → invoice chain"),
    ])

    add_h2(doc, "20 blog post titles, ranked by ROI")
    add_table(
        doc,
        ["#", "Title", "Target keyword", "Words", "CTA"],
        [
            ["1",  "SS 304 vs 316 for Submersible Pump Motor Tubes: 2026 Engineer's Decision Guide", "ss 304 vs 316 submersible pump", "2,200", "MTC sample PDF"],
            ["2",  "ASTM A312 vs A213 vs A268: Which Stainless Tube Standard Your Project Needs", "astm a312 vs a213", "2,000", "Spec sheet bundle"],
            ["3",  "EN 10204 Type 3.1 Mill Test Certificate Explained (with Annotated Sample)", "en 10204 3.1 mtc", "1,800", "Download annotated MTC"],
            ["4",  "4″ vs 6″ vs 8″ Submersible Motor Tube Sizing Chart (HP / GPM / Depth)", "submersible motor tube sizing", "2,400", "Sizing calculator"],
            ["5",  "Stage II Vapour Recovery Decommissioning: Which Aluminum Tubing You Keep", "stage ii vapor recovery decommissioning", "2,200", "RFQ + EPA checklist"],
            ["6",  "Al 6061 vs 6063 vs 5052 for Petrol Vapour Lines: Corrosion + Formability Compared", "al 6061 vs 6063 vs 5052", "1,900", "Material data PDF"],
            ["7",  "Sourcing Stainless Steel Pipes from India: 12-Point Audit Before You Sign", "sourcing stainless steel from india", "2,500", "Supplier audit checklist"],
            ["8",  "UK PPC Vapour Recovery Compliance: Material Specs for New Service Stations", "uk ppc vapour recovery", "1,800", "RFQ + UK spec sheet"],
            ["9",  "Higher HP Deep-Well Motor Tubes (15–50 HP): Long-Tube Manufacturing Matters", "high hp submersible motor tube", "1,600", "Plant tour + RFQ"],
            ["10", "How to Spot a Fake Mill Test Certificate (5 Red Flags Buyers Miss)", "fake mill test certificate", "1,700", "MTC verification call"],
            ["11", "SS Pressure Pump Sleeve Spec Sheet: NRV Pipe Dimensions & Tolerances", "ss pressure pump sleeve", "1,500", "Spec sheet bundle"],
            ["12", "Section 232 Tariffs & CBAM 2026: Real Landed Cost of Indian SS Pipe in US/UK", "section 232 stainless steel india 2026", "2,000", "Landed-cost calculator"],
            ["13", "Pickling vs Passivation vs Electropolishing: What Your Pump Tube ID Finish Should Be", "passivation vs pickling stainless tube", "1,800", "Finish comparison PDF"],
            ["14", "Pitting Corrosion in Coastal Submersible Pumps: 304L vs 316L Field Data", "pitting corrosion submersible pump", "1,600", "Case study PDF"],
            ["15", "ISO 9001:2015 + IS:6761 Combined Audit: What US/UK Importers Should Check", "iso 9001 is 6761 audit", "1,400", "Audit checklist"],
            ["16", "Wall Thickness & Collapse Pressure for Submersible Motor Tubes (Working Formulas)", "submersible motor tube wall thickness", "1,700", "Calculator + tables"],
            ["17", "EN 754 vs ASTM B210: Cross-Reference Table for Aluminum Drawn Tubing Buyers", "en 754 vs astm b210", "1,300", "Cross-ref PDF"],
            ["18", "CIF vs FOB for Stainless Pipe Shipments to US/UK: Hidden Cost Breakdown", "cif vs fob stainless steel pipe india", "1,500", "Incoterm cheatsheet"],
            ["19", "Welded vs Seamless Motor Tubes: When UL/NSF Pump Builders Accept Each", "welded vs seamless motor tube", "1,600", "UL/NSF reference"],
            ["20", "Replacing Copper Vapour Lines With Aluminum: 10-Year TCO for Fuel Retailers", "aluminum vapor recovery tco", "1,500", "TCO calculator"],
        ],
    )

    add_h2(doc, "10 lead magnets, ranked by impression-driving potential")
    add_numbered(doc, [
        "Pump-Tube Sizing Calculator (interactive JS, embeddable — every embed = backlink)",
        "EN 10204 3.1 MTC Annotated Sample PDF (real Apurvi MTC with arrows)",
        "ASTM A312 vs A213 vs A268 Interactive Comparison Tool",
        "Gas-Station Vapour Recovery Compliance Map (US EPA Stage II by state + UK PPC)",
        "Indian SS Pipe Landed-Cost Calculator (Section 232 + CBAM baked in)",
        "Submersible Motor Tube Spec Sheet Bundle (8 PDFs, gated by email)",
        "12-Point Indian Supplier Audit Checklist (Word doc)",
        "Wall Thickness / Collapse Pressure Workbook (Excel)",
        "Pitting Corrosion Field Study — Apurvi's original chloride exposure data (12 pp whitepaper)",
        "Cross-Reference Card: EN 10296 ↔ ASTM A312 ↔ IS:6761 ↔ JIS G3463 (printable A3)",
    ])

    add_page_break(doc)

    # 4. AUTHORITY & LINKS
    add_h1(doc, "4. Authority & Link Building")
    add_para(doc, "Sliding scale: 5–10 DA 60+ links = bare-minimum signal; 20–30 links = long-tail ranking shifts; 50–80 links = mid-tail movement.")

    add_h2(doc, "Tier 1 — Must-do directories (Week 1)")
    add_table(
        doc,
        ["#", "Directory", "Country", "DA", "Type", "Why"],
        [
            ["1",  "ThomasNet", "US", "86", "Free + paid premium", "CRITICAL — #1 US industrial buyer search"],
            ["2",  "GlobalSpec / Engineering360", "US/Global", "85", "Free supplier profile", "IEEE-engineer audience"],
            ["3",  "Kompass UK", "UK", "79", "Free + Booster", "Indexed in 26 languages"],
            ["4",  "Kompass Global / US", "US", "79", "Free company page", "Same authority, US subdir"],
            ["5",  "Europages", "EU", "77", "Free + premium", "2.6M companies, 6M monthly searches"],
            ["6",  "Applegate Marketplace UK", "UK", "64", "Free + RFQ matching", "UK's longest-established"],
            ["7",  "DirectIndustry (VirtualExpo)", "EU/Global", "72", "Paid", "Product page indexing"],
            ["8",  "Manta", "US", "72", "Free claim", "US small/mid B2B"],
            ["9",  "D&B Hoovers / Dun & Bradstreet", "US", "91", "Free DUNS + paid", "Required for US gov contracts"],
            ["10", "MFG.com", "US", "60", "Free + RFQ bidding", "Active RFQs for SS shops"],
            ["11", "IndustryNet", "US", "58", "Free", "US-focused"],
            ["12", "MacRAE'S Blue Book", "US", "60", "Free basic", "Long-running US industrial"],
            ["13", "Yell.com", "UK", "76", "Free + paid", "UK local + B2B trust"],
            ["14", "AZoM Materials", "Global", "73", "Supplier profile", "Engineering audience"],
            ["15", "Stainless Steel World Directory", "NL/Global", "55", "Paid niche", "Direct keyword relevance"],
        ],
    )

    add_h2(doc, "Tier 2 — 20 industry publications to pitch")
    add_bullets(doc, [
        ("Pumps & Systems Magazine (US)", "pumpsandsystems.com/editorial — technical contributed articles"),
        ("WaterWorld Magazine (US)", "waterworld.com/magazine — tech articles + supplier insights"),
        ("Flow Control Magazine (US)", "flowcontrolnetwork.com — guest tech articles"),
        ("Process Industry Forum (UK)", "processindustryforum.com — guest posts, supplier news"),
        ("The Engineer (UK)", "theengineer.co.uk — news + opinion"),
        ("Eureka Magazine (UK)", "eurekamagazine.co.uk — design engineering"),
        ("Petrol Plaza (EU/UK)", "petrolplaza.com — industry interviews"),
        ("CSP Daily News (US)", "cspdailynews.com/fuels — industry features"),
        ("Fuels Market News (US)", "fuelsmarketnews.com — op-eds, supplier features"),
        ("NACE / AMPP Materials Performance", "ampp.org/publications — technical articles"),
        ("Hydrocarbon Engineering (UK)", "hydrocarbonengineering.com — tech features"),
        ("Stainless Steel World (NL)", "stainless-steel-world.net — articles + directory"),
        ("World Pumps (UK)", "worldpumps.com — tech features"),
        ("Modern Pumping Today (US)", "modernpumpingtoday.com — guest articles"),
        ("Empowering Pumps & Equipment (US)", "empoweringpumps.com — thought leadership"),
        ("Engineering & Technology / IET (UK)", "eandt.theiet.org — features"),
        ("Manufacturing Today (UK)", "manufacturing-today.com — supplier features"),
        ("The Manufacturer (UK)", "themanufacturer.com — editorials"),
        ("Process Engineering (UK)", "processengineering.co.uk — technical articles"),
        ("Thomas Insights (US)", "blog.thomasnet.com — sourcing thought leadership"),
    ])

    add_callout(
        doc,
        "Sample pitch (use for ALL)",
        "Subject: Contributed article — 'How Indian Tube Manufacturers Meet ASTM A312 for US Pump OEMs'. "
        "Body: (1) intro Apurvi credentials — ISO 9001, EN 10204 3.1 MTC, 25–200 HP specialization; (2) three article angles; "
        "(3) word count + delivery date; (4) author bio linking to apurviind.com.",
    )

    add_h2(doc, "Tier 3 — Reddit / Quora / Stack Exchange")
    add_callout(
        doc,
        "Reddit rule of engagement",
        "10 helpful comments before any link drop. Always disclose 'I work at a stainless tube manufacturer in India'. "
        "Tailor every response. Throwaway karma-building period of 2 weeks first.",
    )
    add_bullets(doc, [
        ("r/AskEngineers (1.4M)", "304L vs 316L for pump applications, PREN explained"),
        ("r/Welding (350K)", "TIG welding ASTM A312 tubes, butt-weld vs socket-weld"),
        ("r/MetalManufacturing (25K)", "Cold-drawn vs hot-rolled tube manufacturing"),
        ("r/Manufacturing (200K)", "India manufacturing capabilities deep-dives"),
        ("r/Plumbing (800K)", "Stainless vs galvanized for industrial water lines"),
        ("r/Procurement (60K)", "What buyers should ask Indian suppliers"),
        ("r/SmallBusiness (2M)", "Case study: 'How we landed a US distributor'"),
        ("r/IndianBusiness", "Export documentation, MTC compliance, FOB pricing"),
    ])

    add_h2(doc, "Tier 4 — HARO / Featured.com / Qwoted (10 ready angles)")
    add_numbered(doc, [
        "China+1 winners — concrete US buyer shift to India",
        "Aluminum supply chain — 6061/6063 for US fuel infrastructure",
        "EV charging vs vapor recovery demand",
        "Indian MSME export success story",
        "EN 10204 3.1 MTC explainer",
        "Welding certifications India vs Europe",
        "Sustainable / recycled stainless steel content",
        "Lead-time benchmarks: India 4–6 wk vs China 8–12 wk",
        "Family-owned manufacturer story (founder Viren Gathani)",
        "Section 232 tariff impact on industrial buyers",
    ])
    add_bullets(doc, [
        ("Featured.com", "featured.com — daily expert queries"),
        ("Qwoted", "qwoted.com — highest HARO-alt conversion"),
        ("Source of Sources", "sourceofsources.com — Peter Shankman successor"),
        ("ResponseSource UK", "responsesource.com — UK journalist requests"),
        ("#JournoRequest on X", "Fastest response window"),
    ])

    add_h2(doc, "Tier 5 — LinkedIn (founder + company)")
    add_bullets(doc, [
        ("Cadence", "5 founder posts/week + 1 LinkedIn Article every 2 weeks"),
        ("LinkedIn Newsletter", "'Tubing Talk: Inside Industrial Steel Manufacturing' — biweekly, 1,000–1,500 words"),
        ("Sales Navigator", "30 connect requests/day to US/UK Procurement Managers, Plant Engineers, Pump OEM Sourcing Heads"),
        ("Showcase Pages", "One per product line (Motor Tubes / Vapour Tubing / Pump Sleeves) — each separately indexable"),
        ("Hashtag mix", "3 broad (#manufacturing #stainlesssteel #B2B) + 3 niche (#ASTMA312 #vaporrecovery #pumpcomponents) + 3 geo (#USmanufacturing #UKindustry)"),
    ])

    add_h2(doc, "Tier 6 — Strategic partnerships (5–8 in 30 days)")
    add_bullets(doc, [
        "US/UK pump OEM channel partners (Grundfos, Franklin Electric, Pentair, Xylem distributors)",
        "Gas station equipment distributors (OPW, Gilbarco, Fairbanks Pumps UK)",
        "Water treatment integrators (Evoqua, Veolia US, Severn Trent suppliers UK)",
        "Pump rebuild / repair shops (regional US/UK)",
        "Trade associations (BSSA UK, Hydraulic Institute US, BPMA UK)",
        "Procurement marketplaces (Scoutbee, Tealbook supplier discovery)",
    ])

    add_page_break(doc)

    # 5. PAID LAYER
    add_h1(doc, "5. Paid Campaign Layer — $8–10K to Bridge the Impression Gap")
    add_table(
        doc,
        ["Channel", "Budget (USD/mo)", "Impressions", "Why"],
        [
            ["Google Search (US+UK, high-intent only)", "$2,500", "~40K", "Highest-intent B2B"],
            ["Bing Ads (US, Copilot inventory)", "$1,500", "~50K", "BEST ROI in 2026 — Copilot is 35-40% of US B2B; CPC 33-40% cheaper"],
            ["LinkedIn Ads (procurement titles)", "$2,500", "~50K", "Highest B2B precision"],
            ["Reddit Ads (r/Welding, r/AskEngineers)", "$800", "~80K", "Cheap, niche-targeted"],
            ["YouTube TrueView (factory tour 15s)", "$1,000", "~80K", "Brand storytelling"],
            ["Google Display (industrial whitelist)", "$400", "~100K", "Cheap vanity impressions"],
            ["Quora Ads", "$300", "~40K", "Long-tail buyer questions"],
            ["TOTAL", "$9,000", "~440K", ""],
        ],
    )
    add_callout(
        doc,
        "Lean alternative",
        "$3,500/mo (GDN + Reddit + Bing only) = ~300K vanity impressions. Use only if budget is tight and lead quality is acceptable.",
    )

    add_h2(doc, "Paid risk caveats")
    add_bullets(doc, [
        "Industrial/manufacturing flag — submit business verification on Day 1",
        "Geo-exclude India / Pakistan / Bangladesh from US/UK campaigns",
        "Frequency cap on Display + YouTube — prevent CPM inflation",
        "LinkedIn CPM up 15% YoY — lock budget early-month",
    ])

    add_page_break(doc)

    # 6. COMPETITOR TEARDOWN
    add_h1(doc, "6. Competitor Teardown — Who to Leapfrog")
    add_h2(doc, "US Top 10")
    add_table(
        doc,
        ["#", "Domain", "HQ", "Type", "DA", "Strength"],
        [
            ["1",  "pennstainless.com", "Quakertown, PA", "Mill-tier distributor", "50–55", "Deep technical content; ranks #1 on most A312 long-tails"],
            ["2",  "continentalsteel.com", "Fort Lauderdale, FL", "Distributor", "55–60", "Largest content moat"],
            ["3",  "plymouth.com (Plymouth Tube)", "Warrenville, IL", "Manufacturer (5 US mills)", "50–55", "True manufacturer credentials"],
            ["4",  "twmetals.com", "Exton, PA", "Distributor", "55", "Strong rich-snippet game"],
            ["5",  "usmetals.com", "Houston, TX", "Distributor", "40–45", "Oil & gas vertical"],
            ["6",  "atlassteel.com", "Twinsburg, OH", "Manufacturer", "35–40", "Owns A268-specific ranking"],
            ["7",  "eagletube.com", "Trevose, PA", "Mfr/processor", "35", "Process-led pages"],
            ["8",  "worldwidepipe.com", "Houston, TX", "Distributor", "35–40", "OCTG link equity"],
            ["9",  "stainlessandalloy.com", "Statham, GA", "Distributor", "30", "Lean, fast"],
            ["10", "opwglobal.com", "Hodgkins, IL", "Mfr", "65–70", "Owns EPA Stage I/II vapor SERP"],
        ],
    )
    add_h2(doc, "UK Top 10")
    add_table(
        doc,
        ["#", "Domain", "HQ", "Type", "DA", "Strength"],
        [
            ["1",  "aalco.co.uk", "Surrey", "Multi-metal stockist (18 centres)", "50–55", "Owns 'stainless steel tube UK'"],
            ["2",  "marcegaglia.co.uk", "West Midlands", "Manufacturer (largest UK)", "55", "Manufacturer + scale"],
            ["3",  "smithmetal.com", "Bedfordshire", "Stockholder + UKAS lab", "45–50", "UKAS lab = compliance trust"],
            ["4",  "materials.sandvik (UK)", "Sheffield", "Manufacturer", "80+", "Brand authority"],
            ["5",  "outokumpu.com (UK)", "Sheffield", "Manufacturer", "65", "Coil/long products"],
            ["6",  "steelexpress.co.uk", "Wolverhampton", "Stockist", "35–40", "Wins seamless tube UK"],
            ["7",  "themetalstore.co.uk", "Manchester", "Online retailer", "35–40", "E-commerce UX"],
            ["8",  "stainlessandaluminium.co.uk", "Tipton", "Stockist", "25–30", "Page-per-grade"],
            ["9",  "nero.co.uk", "Bristol", "Pipe fitting stockist", "30–35", "Same-day dispatch"],
            ["10", "bssa.org.uk", "Sheffield", "Trade body", "55", "Top-3 for EN 10296-2 — must-target backlink"],
        ],
    )

    add_h2(doc, "Content gaps Apurvi can win (zero/weak competitor coverage)")
    add_numbered(doc, [
        "EN 10204 3.1 MTC walkthrough for US buyers — competitors mention it; none explain it",
        "ASTM A268 vs A269 vs A312 full comparison matrix — nobody has it",
        "India vs China stainless steel sourcing 2026 — zero competitor content",
        "How US procurement qualifies an Indian SS tube manufacturer — buyer-journey gap",
        "Submersible pump motor shell material selection (304 vs 316 vs duplex) — pump OEMs don't publish; tube suppliers don't either",
        "Aluminum vapor-recovery rigid tubing — Stage I balance retrofit (OPW sells systems, nobody owns the tubing replacement-part SERP)",
        "EN 10296-2 vs ASTM A554 dimensional tolerance comparison",
        "High HP submersible pump tube — 50/75/100 HP shell sizing (pure long-tail Apurvi can sweep)",
        "Section 232 + duty-drawback for Indian exporters — timely, pulls procurement buyers",
        "NRV (non-return valve) sizing chart for submersible pump risers — English-language version doesn't exist well",
    ])

    add_h2(doc, "The one weakness every Indian exporter has — and how Apurvi sidesteps it")
    add_callout(
        doc,
        "Common weakness",
        "Indian exporter sites read like Indian-domestic sites: no US/UK phone, no local addresses, no US testimonials, MTCs buried, "
        "IndiaMART-only backlinks, WhatsApp-only contact, INR pricing.",
    )
    add_numbered(doc, [
        "Get a US 1-800 + UK 0800 forwarding number",
        "Set up US Inc / UK Ltd shell or registered-agent address (drop-shipping/forwarder OK)",
        "Build /us/ and /uk/ subfolders with hreflang (NOT subdomains)",
        "Pin SGS / Bureau Veritas / TUV inspection logos in the hero of every product page",
        "Host 2–3 US case studies as soon as orders close — even small ones",
        "Embed sample MTC PDF (anonymized) on every grade page",
        "Get listed on Thomasnet — single highest-ROI US B2B backlink",
        "Avoid 'manufacturer in India' in /us/ titles — write 'ASTM A312 TP304 supplier — MTC 3.1 certified, USA delivery'",
        "Reference Section 232 reality on every US page — '~20–30% landed-cost advantage vs Chinese-origin'",
        "Replace WhatsApp-only with Calendly link for sales calls",
    ])

    add_page_break(doc)

    # 7. 30-DAY CALENDAR
    add_h1(doc, "7. The 30-Day Calendar (Day-by-Day)")

    add_h2(doc, "Week 1 (29 Jun – 5 Jul) — Foundation + Indexing")
    add_table(
        doc,
        ["Day", "Task"],
        [
            ["1", "Ship on-site code: hreflang fix, robots.txt update, llms-full.txt, US/UK schema (DONE — see tech brief)"],
            ["1", "Submit apurviind.com to Bing Webmaster Tools. Verify Domain in GSC."],
            ["1", "Create skeleton /us/ and /uk/ folders (homepage + products + contact)"],
            ["2", "Submit sitemaps for /us/ and /uk/ to GSC + Bing. Ping IndexNow."],
            ["2", "Apply for DUNS (D&B Hoovers — 30-day issuance)"],
            ["2", "Sign up US 1-800 + UK 0800 forwarding (Twilio / OpenPhone / Voipfone)"],
            ["3", "List on ThomasNet, GlobalSpec, Kompass US, Kompass UK, Europages"],
            ["3", "Launch paid accounts: Google, Bing, LinkedIn, Reddit, YouTube. Submit business verification."],
            ["4", "Publish blog #1: SS 304 vs 316 for Submersible Pump Motor Tubes"],
            ["4", "List on Applegate, DirectIndustry, Manta, MFG.com, AZoM"],
            ["5", "Publish blog #2: ASTM A312 vs A213 vs A268"],
            ["5", "Author bios live for Viren Gathani + senior engineer (E-E-A-T)"],
            ["6", "Publish blog #3: EN 10204 Type 3.1 MTC Explained + Lead Magnet #2 ships (MTC PDF)"],
            ["6", "LinkedIn Article #1 published by founder. Cross-post to Medium with canonical."],
            ["7", "YouTube short #1: factory tour (60s, ASTM mark visible). Embed in homepage."],
            ["7", "First Reddit answers in r/AskEngineers (10 helpful comments, zero links)"],
        ],
    )

    add_h2(doc, "Week 2 (6 Jul – 12 Jul) — Vapour Recovery + Compliance")
    add_table(
        doc,
        ["Day", "Task"],
        [
            ["8",  "Publish blog #5: Stage II Vapour Recovery Decommissioning + LinkedIn"],
            ["8",  "Launch paid campaigns LIVE — Google + Bing + Reddit + LinkedIn + Display"],
            ["9",  "Publish blog #6: Al 6061 vs 6063 vs 5052 for Petrol Vapour Lines"],
            ["10", "Publish Pillar B (3,500 words): Aluminum Vapour Recovery Tubing Compliance Guide"],
            ["10", "Pitch posts #5 + #8 to NACS Magazine, Convenience Store News, Petrol Plaza"],
            ["11", "Publish blog #8: UK PPC Vapour Recovery Compliance"],
            ["11", "Lead Magnet #4 ships: Vapour Recovery Compliance Map (US state + UK region)"],
            ["12", "Reddit answers across r/Welding, r/MetalManufacturing"],
            ["13", "LinkedIn Article #2 by founder: China+1 — What US Procurement Is Doing in 2026"],
            ["14", "YouTube short #2: Aluminum vapour tubing close-up + EXIF metadata"],
        ],
    )

    add_h2(doc, "Week 3 (13 Jul – 19 Jul) — Trust + Sourcing")
    add_table(
        doc,
        ["Day", "Task"],
        [
            ["15", "Publish blog #7 (Pillar C kickoff): Sourcing SS Pipes from India — 12-Point Audit"],
            ["16", "Publish blog #10: How to Spot a Fake MTC (5 Red Flags)"],
            ["17", "Publish blog #12: Section 232 + CBAM 2026 — Real Landed Cost"],
            ["17", "Lead Magnet #7 ships: 12-Point Supplier Audit Checklist"],
            ["18", "Lead Magnet #5 ships: Indian SS Pipe Landed-Cost Calculator"],
            ["19", "Publish Pillar C (3,500 words)"],
            ["20", "First HARO/Featured.com pitches (5 angles)"],
            ["21", "LinkedIn Newsletter #1 launches: 'Tubing Talk' inaugural issue"],
        ],
    )

    add_h2(doc, "Week 4 (20 Jul – 28 Jul) — Spec Depth + Long-Tail Sweep")
    add_table(
        doc,
        ["Day", "Task"],
        [
            ["22", "Publish blog #4: 4″/6″/8″ Sizing Chart + Lead Magnet #1 ships (Pump-Tube Sizing Calculator)"],
            ["22", "Outreach to 30 pump distributors offering Calculator embed widget"],
            ["23", "Publish blog #9: Higher HP (15–50 HP) Deep-Well Motor Tubes"],
            ["24", "Publish blog #11 + #13 spec sheet day"],
            ["25", "Publish Pillar A (3,500 words)"],
            ["26", "YouTube shorts #3 + #4: MTC walkthrough + 4″/6″/8″ animation"],
            ["27", "Audit AI Overview appearances using Otterly/Profound. Tweak underperforming posts."],
            ["28", "LinkedIn Article #3 by founder + LinkedIn Newsletter #2"],
            ["29", "Final HARO/Featured submissions"],
            ["30", "Generate monthly report: GSC impressions split US/UK + Bing AI Performance + RFQ count"],
        ],
    )

    add_h2(doc, "Daily operations (every day, ~30 min)")
    add_bullets(doc, [
        "10 LinkedIn comments on procurement/engineering posts",
        "2 Reddit helpful answers (karma building)",
        "1 Quora answer",
        "Track citations in ChatGPT / Perplexity / AIO via Otterly or Profound",
    ])

    add_page_break(doc)

    # 8. KPI DASHBOARD
    add_h1(doc, "8. KPI Dashboard (track weekly)")
    add_table(
        doc,
        ["Metric", "Vanity?", "Target Month 1"],
        [
            ["Total search impressions (GSC US + UK split)", "Partial", "250K–400K"],
            ["AI Overview / Bing Copilot citations", "No", "8–15 (Bing WMT AI Performance)"],
            ["Organic clicks (US + UK)", "No", "4,000–8,000"],
            ["Qualified RFQs from US/UK IPs", "PRIMARY", "30"],
            ["Cost per qualified RFQ", "No", "< $400"],
            ["Brand search volume (WoW growth)", "No", "+30%"],
            ["Referring domains (US/UK origin)", "No", "+15 new"],
            ["LinkedIn Article impressions", "Partial", "30K"],
            ["YouTube view-through (15s+)", "Partial", "8K"],
            ["Newsletter / MTC sample downloads", "No", "200"],
            ["Pump-Tube Sizing Calculator embeds (3rd-party)", "No", "5"],
        ],
    )
    add_callout(
        doc,
        "Cardinal rule",
        "400K impressions with < 20 qualified RFQs = failed campaign. Pivot if that ratio breaks at week 2.",
    )

    add_page_break(doc)

    # 9. OFF-SITE ACTION LIST
    add_h1(doc, "9. Off-Site Action List (Owner: SEO + Marketing Team)")
    add_table(
        doc,
        ["Action", "Day", "Critical?"],
        [
            ["Bing Webmaster Tools verification + sitemap", "Day 1", "CRITICAL"],
            ["Google Search Console — verify Domain property + sitemap", "Day 1", "CRITICAL"],
            ["Apply for D&B DUNS number", "Day 2", "CRITICAL (30-day issuance)"],
            ["US 1-800 + UK 0800 forwarding numbers", "Day 2", "CRITICAL"],
            ["ThomasNet listing (business.thomasnet.com/get-listed-on-thomasnet)", "Day 3", "CRITICAL"],
            ["Kompass US + UK listings", "Day 3", "HIGH"],
            ["Europages + Applegate UK listings", "Day 4", "HIGH"],
            ["GlobalSpec / Engineering360 supplier profile", "Day 4", "HIGH"],
            ["Launch Google + Bing + LinkedIn + Reddit Ads accounts; business verification", "Day 3", "CRITICAL"],
            ["ISO 9001 + SGS / TUV / BV logos sourced and added to product pages", "Day 5", "CRITICAL US trust"],
            ["Founder LinkedIn refresh (Viren Gathani — photo, banner, headline)", "Day 5", "HIGH"],
            ["Sign up for Featured.com / Qwoted / SOS / ResponseSource UK", "Day 6", "HIGH"],
            ["BSSA membership application (UK trade body)", "Week 2", "MEDIUM"],
            ["LinkedIn Sales Navigator subscription", "Week 1", "HIGH"],
            ["2–3 US/UK customer testimonials sourced + photos", "Week 2", "CRITICAL"],
            ["Calendly for sales calls (replace WhatsApp-only)", "Day 5", "CRITICAL"],
            ["ISSF + ISSDA + BSSA membership listings", "Week 2", "MEDIUM"],
            ["3 monthly LinkedIn Articles by founder", "Ongoing", "HIGH"],
        ],
    )

    # 10. RISK REGISTER
    add_h1(doc, "10. Risk Register")
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ["Indexing delays (new pages stuck 14–30 days)", "IndexNow ping every publish + internal linking from existing 16 pages + GSC URL Inspection API"],
            ["Paid account suspension (industrial flagged)", "Submit business verification Day 1; clean landing pages; no exaggerated claims"],
            ["AI Overview click cannibalization (-39% CTR at pos 2)", "Optimize for BEING cited in AI Overviews — structured data, factual specs, comparison tables"],
            ["Manual penalty from spammy directory listings", "Tier 1 directories only (no bulk submissions)"],
            ["Google spam updates (Aug/Mar cadence)", "Stick to deep technical content; avoid thin AI-generated pages"],
            ["LinkedIn CPM inflation", "Lock budget early-month, watch frequency caps"],
            ["Hreflang misconfiguration", "Self-reference canonical on every locale; bidirectional alternates"],
            ["Wrong-country impressions", "Geo-exclude India / Pakistan / Bangladesh in US/UK ad sets"],
            ["Vanity 400K with no leads", "Track qualified RFQs as PRIMARY metric — pivot at week 2 if ratio breaks"],
        ],
    )

    # 11. MONTH 2+
    add_h1(doc, "11. Month 2 and Beyond")
    add_bullets(doc, [
        "600K–1M impressions (paid amortizes, organic compounds)",
        "80+ qualified RFQs",
        "5+ podcast appearances on industrial/procurement shows",
        "First US/UK case study published",
        "Pump-Tube Sizing Calculator → 30+ 3rd-party embeds",
        "BSSA / Hydraulic Institute / BPMA memberships finalised",
        "By month 6: top-10 ranking on Google US for at least 5 mid-tail keywords + weekly Perplexity/ChatGPT citations for 'non-China stainless tube supplier' type queries",
    ])

    doc.save("/Users/sayujpillai/Desktop/apurvi/Apurviind.com/Apurvi-SEO-Team-Brief.docx")
    print("SEO doc saved.")


# =============================================================================
#                     DOC B — TECH TEAM BRIEF
# =============================================================================

def build_tech_doc():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    add_cover(
        doc,
        title="US + UK SEO — Technical Implementation Brief",
        subtitle="Hreflang, Schema, AEO, Core Web Vitals & Lead Magnet Engineering",
        audience="Tech / Development Team",
    )
    add_page_break(doc)

    # 1. SCOPE
    add_h1(doc, "1. Scope & Context")
    add_para(
        doc,
        "Apurvi Industries is targeting US + UK SEO expansion to hit 400,000 search impressions in 30 days. "
        "The SEO and marketing team will produce content, run paid campaigns, and pursue backlinks (see SEO Brief). "
        "This document is the tech team's complete implementation list: what to ship, in what order, with code samples.",
    )
    add_callout(
        doc,
        "Baseline (good news)",
        "The current site already has solid foundations: 13 HTML pages with full schema graph (Organization + LocalBusiness + WebSite + FAQPage + Speakable), "
        "robots.txt allowing GPTBot + ClaudeBot, llms.txt, IndexNow key file, sitemap.xml with 15 URLs, BreadcrumbList on every internal page, "
        "Article schema on each blog post.",
    )
    add_callout(
        doc,
        "What's broken for US/UK ranking",
        "(1) hreflang en-IN/en/x-default ALL point to the same homepage — Google ignores this. (2) geo.region=IN-GJ on every page (including service-areas) — signals 'Indian site' to Google. "
        "(3) Titles all say 'India' / 'Ahmedabad' — won't rank in US/UK Google. (4) UK missing from areaServed in schema. (5) No US distributor address, no US phone, no US testimonials. "
        "(6) Only 3 blog posts — too thin for US/UK competition.",
    )

    add_page_break(doc)

    # 2. ALREADY SHIPPED
    add_h1(doc, "2. What's Already Shipped (29 June 2026)")
    add_bullets(doc, [
        ("robots.txt", "Added explicit allow rules for OAI-SearchBot, ChatGPT-User, PerplexityBot, Perplexity-User, Claude-SearchBot, Claude-User, Google-Extended, Applebot-Extended, CCBot, Bytespider, Amazonbot, Diffbot, YouBot. Disallowed /admin.php, /send-inquiry.php, /private/."),
        ("llms.txt", "Restructured to lead with export markets + US/UK buyer-angled FAQ block (8 Q's). Tightened first 200 tokens for AEO extraction."),
        ("llms-full.txt", "NEW file. Single Markdown knowledge dump of products + standards + compliance + exports + 2026 tariff context. Probed by Cursor, Claude Code, GitHub Copilot, Cline, Aider, Bing Copilot."),
        ("index.html schema", "Extended ContactPoint.areaServed and LocalBusiness.areaServed to include United Kingdom. Added GBP to currenciesAccepted. Added Wikidata sameAs entity links to all country references (improves AI entity resolution)."),
        ("INTERNATIONAL-SEO-PLAN.md", "Master plan committed to repo root."),
    ])

    add_page_break(doc)

    # 3. HREFLANG
    add_h1(doc, "3. Hreflang Strategy (Priority 1 — ship this first)")
    add_para(
        doc,
        "Recommendation: en-US / en-GB at SUBFOLDERS (NOT ccTLDs, NOT subdomains). Subfolders consolidate link equity, "
        "are cheaper to maintain, and Google accepts them as a valid geo-targeting signal alongside hreflang.",
    )

    add_h2(doc, "Target URL structure")
    add_code_block(doc,
        "https://apurviind.com/          ← India default, en-IN\n"
        "https://apurviind.com/us/       ← United States, en-US\n"
        "https://apurviind.com/uk/       ← United Kingdom, en-GB\n"
        "\n"
        "Cloned for each region (with locale-appropriate copy/spelling):\n"
        "  /us/products    /uk/products\n"
        "  /us/industries  /uk/industries\n"
        "  /us/contact     /uk/contact\n"
        "  /us/company     /uk/company"
    )

    add_h2(doc, "Hreflang link tags — drop into <head> of every regional homepage")
    add_code_block(doc,
        '<link rel="alternate" hreflang="en-in" href="https://apurviind.com/" />\n'
        '<link rel="alternate" hreflang="en-us" href="https://apurviind.com/us/" />\n'
        '<link rel="alternate" hreflang="en-gb" href="https://apurviind.com/uk/" />\n'
        '<link rel="alternate" hreflang="x-default" href="https://apurviind.com/" />\n'
        '<link rel="canonical" href="https://apurviind.com/us/" />   <!-- self-reference on US page -->'
    )
    add_callout(
        doc,
        "Critical rule",
        "Every regional page's canonical must SELF-REFERENCE (point to itself), not back to the India version. "
        "Hreflang must be BIDIRECTIONAL — every alternate page declares all the others. "
        "Repeat this block on /us/products, /us/industries, etc. (with corresponding /uk/ counterparts).",
    )

    add_h2(doc, "Sitemap-level hreflang (additional reinforcement)")
    add_code_block(doc,
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"\n'
        '        xmlns:xhtml="http://www.w3.org/1999/xhtml">\n'
        '  <url>\n'
        '    <loc>https://apurviind.com/us/</loc>\n'
        '    <xhtml:link rel="alternate" hreflang="en-us" href="https://apurviind.com/us/" />\n'
        '    <xhtml:link rel="alternate" hreflang="en-gb" href="https://apurviind.com/uk/" />\n'
        '    <xhtml:link rel="alternate" hreflang="en-in" href="https://apurviind.com/" />\n'
        '    <xhtml:link rel="alternate" hreflang="x-default" href="https://apurviind.com/" />\n'
        '  </url>\n'
        '</urlset>'
    )

    add_h2(doc, "Geo meta tags — fix conflict")
    add_bullets(doc, [
        ("Current state", "geo.region=IN-GJ on EVERY page including /service-areas — signals 'Indian site' to crawlers."),
        ("Fix", "Remove geo.region/geo.position/geo.placename/ICBM tags entirely from /us/* and /uk/* pages. (They add zero SEO value in 2026 — hreflang + schema areaServed do the work.) Leave them on the India root for honesty."),
    ])

    add_page_break(doc)

    # 4. SCHEMA
    add_h1(doc, "4. Schema Extensions for International Targeting")
    add_para(
        doc,
        "Existing LocalBusiness schema must stay scoped to the India HQ/plant (truthful). A parallel Organization node with areaServed covering US + UK "
        "carries the international targeting weight. Service nodes get country-specific areaServed.",
    )

    add_h2(doc, "Drop this @graph block in /us/index.html <head>")
    add_code_block(doc,
        '<script type="application/ld+json">\n'
        '{\n'
        '  "@context": "https://schema.org",\n'
        '  "@graph": [\n'
        '    {\n'
        '      "@type": "Organization",\n'
        '      "@id": "https://apurviind.com/#organization",\n'
        '      "name": "Apurvi Industries Pvt. Ltd.",\n'
        '      "url": "https://apurviind.com/us/",\n'
        '      "logo": "https://apurviind.com/assets/logo.png",\n'
        '      "sameAs": ["https://www.linkedin.com/company/apurvi-ind/"],\n'
        '      "founder": {"@type": "Person", "name": "Viren Gathani"},\n'
        '      "areaServed": [\n'
        '        {"@type": "Country", "name": "United States", "sameAs": "https://www.wikidata.org/wiki/Q30"},\n'
        '        {"@type": "Country", "name": "United Kingdom", "sameAs": "https://www.wikidata.org/wiki/Q145"},\n'
        '        {"@type": "Country", "name": "Canada"},\n'
        '        {"@type": "Country", "name": "Germany"}\n'
        '      ],\n'
        '      "knowsAbout": ["Stainless Steel Motor Tubes","Submersible Pump Tubing",\n'
        '                     "SS 304","SS 316","ASTM A268","ASTM A312","EN 10296"],\n'
        '      "contactPoint": [\n'
        '        {"@type": "ContactPoint","contactType": "sales","areaServed": "US",\n'
        '         "availableLanguage": "en","email": "info@apurviind.com",\n'
        '         "telephone": "+91-8128664329"}\n'
        '      ]\n'
        '    },\n'
        '    {\n'
        '      "@type": "Service",\n'
        '      "@id": "https://apurviind.com/us/#service-ss-tubes",\n'
        '      "serviceType": "Stainless Steel Motor Tube Manufacturing & Export",\n'
        '      "provider": {"@id": "https://apurviind.com/#organization"},\n'
        '      "areaServed": [\n'
        '        {"@type": "AdministrativeArea","name": "Texas"},\n'
        '        {"@type": "AdministrativeArea","name": "California"},\n'
        '        {"@type": "AdministrativeArea","name": "Illinois"},\n'
        '        {"@type": "AdministrativeArea","name": "Ohio"},\n'
        '        {"@type": "AdministrativeArea","name": "Pennsylvania"}\n'
        '      ]\n'
        '    }\n'
        '  ]\n'
        '}\n'
        '</script>'
    )

    add_h2(doc, "Future: when a real US distributor is signed")
    add_para(doc, "Register them as a separate LocalBusiness node with their REAL US street address. Do NOT use fake US addresses — Google penalises this aggressively.")

    add_page_break(doc)

    # 5. AEO ARTICLE SKELETON
    add_h1(doc, "5. AEO-Ready Article Template (MANDATORY for all new blog posts)")
    add_para(
        doc,
        "Every new post must use this exact HTML/H2/H3 structure. It's the format that wins Google AI Overviews, ChatGPT, Perplexity, and Bing Copilot citations in 2026 "
        "(44.2% of ChatGPT citations come from the first 30% of a page — lead with the direct answer).",
    )
    add_code_block(doc,
        '<article>\n'
        '  <h1>{Specific, question-shaped title under 65 chars}</h1>\n'
        '  <p class="updated">Last updated: 29 June 2026 · 8 min read ·\n'
        '     Author: <a href="/team/viren-gathani">Viren Gathani</a>, Founder</p>\n'
        '\n'
        '  <!-- Direct answer block: 40-75 words, complete and quotable -->\n'
        '  <h2>What is {topic}? (Direct answer)</h2>\n'
        '  <p>The short answer: SS 316L is required for submersible motors in chloride water\n'
        '     above 200 ppm; SS 304L is acceptable below that. Per ASTM A312 and A213,\n'
        '     the 2-3% Mo content in 316L gives ~5x better pitting resistance (PREN 24 vs 19).</p>\n'
        '\n'
        '  <h2>Key specifications at a glance</h2>\n'
        '  <table>\n'
        '    <tr><th>Spec</th><th>SS 304L</th><th>SS 316L</th></tr>\n'
        '    <tr><td>Cr %</td><td>18-20</td><td>16-18</td></tr>\n'
        '    <tr><td>Mo %</td><td>—</td><td>2-3</td></tr>\n'
        '    <tr><td>PREN</td><td>~19</td><td>~24</td></tr>\n'
        '  </table>\n'
        '\n'
        '  <h2>Standards cited</h2>\n'
        '  <p>ASTM A312, A213; EN 10296; IS:6761; EN 10204 Type 3.1.\n'
        '     <a href="https://www.astm.org/a0312-19.html">ASTM A312 specification →</a></p>\n'
        '\n'
        '  <h2>Field data (Apurvi original photos with EXIF + captions)</h2>\n'
        '  <figure><img src="..." alt="Heat ID stamping on SS 316L motor tube"/>\n'
        '          <figcaption>Heat number 2025-AP-3041 ...</figcaption></figure>\n'
        '\n'
        '  <h2>Common questions about {topic}</h2>\n'
        '  <h3>Is SS 316L always better than SS 304L for pump motors?</h3>\n'
        '  <p>{40-60 word answer}</p>\n'
        "  <!-- ... 4-6 more Q and A pairs ... -->\n"
        '\n'
        '  <h2>Engineer\'s checklist before issuing an RFQ</h2>\n'
        '  <ol> ... numbered, AI-quotable ... </ol>\n'
        '\n'
        '  <h2>Sources and further reading</h2>\n'
        '  <ol><li><a href="https://www.astm.org/...">ASTM A312 specification</a></li>...</ol>\n'
        '</article>\n'
        '\n'
        '<!-- Required schema blocks -->\n'
        '<script type="application/ld+json">{Article + FAQPage + BreadcrumbList}</script>'
    )

    add_h2(doc, "Required schema per article")
    add_bullets(doc, [
        ("Article schema", "headline, datePublished, dateModified (CRITICAL — recency is a top-3 AIO ranking signal), author (Person with sameAs to LinkedIn), publisher (Organization @id reference)"),
        ("FAQPage schema", "5–7 Question/Answer pairs, answers 40-60 words. Do NOT over-stuff — Seer 2026 data shows >7 FAQs correlates negatively with citation share"),
        ("BreadcrumbList", "/ > /blog > {post}"),
        ("Image schema", "ImageObject with contentUrl, license, creator, EXIF date"),
    ])

    add_page_break(doc)

    # 6. ROBOTS.TXT
    add_h1(doc, "6. robots.txt — Current State (DEPLOYED)")
    add_code_block(doc,
        "User-agent: *\n"
        "Allow: /\n"
        "Disallow: /wp-admin/\n"
        "Disallow: /wp-login.php\n"
        "Disallow: /private/\n"
        "Disallow: /admin.php\n"
        "Disallow: /send-inquiry.php\n"
        "\n"
        "# AI training crawlers — citation eligibility\n"
        "User-agent: GPTBot\nAllow: /\n\n"
        "User-agent: ClaudeBot\nAllow: /\n\n"
        "User-agent: Google-Extended\nAllow: /\n\n"
        "User-agent: Applebot-Extended\nAllow: /\n\n"
        "User-agent: CCBot\nAllow: /\n\n"
        "User-agent: Bytespider\nAllow: /\n\n"
        "# AI search/retrieval crawlers — CRITICAL for citations\n"
        "User-agent: OAI-SearchBot\nAllow: /\n\n"
        "User-agent: ChatGPT-User\nAllow: /\n\n"
        "User-agent: PerplexityBot\nAllow: /\n\n"
        "User-agent: Perplexity-User\nAllow: /\n\n"
        "User-agent: Claude-SearchBot\nAllow: /\n\n"
        "User-agent: Claude-User\nAllow: /\n\n"
        "User-agent: Amazonbot\nAllow: /\n\n"
        "Sitemap: https://apurviind.com/sitemap.xml"
    )

    add_page_break(doc)

    # 7. CORE WEB VITALS
    add_h1(doc, "7. Core Web Vitals & Performance (Lighthouse 2026)")
    add_callout(
        doc,
        "Top issue in 2026",
        "INP (Interaction to Next Paint) is now the most-failed CWV — 43% of sites fail the 200ms threshold. Lighthouse weights TBT at 30% and LCP at 25%.",
    )

    add_h2(doc, "Wins to implement")
    add_bullets(doc, [
        ("INP / TBT", "Audit JavaScript click handlers. Break any task >50ms with scheduler.yield() or requestIdleCallback(). Avoid bundle-bombs."),
        ("Self-host fonts", 'Use <link rel="preload" href="/assets/fonts/inter-var.woff2" as="font" type="font/woff2" crossorigin>. Subset to Latin-only. font-display: swap. size-adjust in @font-face to prevent CLS.'),
        ("Hero images → AVIF", "Use <picture> with AVIF + WebP fallback. Set explicit width/height attributes to prevent CLS. fetchpriority='high' on LCP image."),
        ("Inline critical CSS", "Inline above-the-fold styles. Defer the rest: <link rel='preload' as='style' onload=\"this.rel='stylesheet'\">."),
        ("LCP > 2.5s rescue", "fetchpriority='high' + <link rel='preload' as='image'>. Drop third-party fonts/scripts from above-the-fold render path."),
        ("CDN", "Cloudflare free tier — HTTP/3 + Brotli compression."),
        ("Mobile tap targets", ">= 48x48px with 8px spacing. Verify 'Get a Quote' buttons on mobile."),
        ("Colour contrast", ">= 4.5:1 for body text. Audit navy/grey on white."),
        ("HTML lang attribute", "<html lang='en-US'> on /us/ pages, lang='en-GB' on /uk/ pages, lang='en-IN' on root. Match hreflang."),
        ("Accessible names", "aria-label on every icon-only button. Alt text descriptive, not keyword-stuffed."),
    ])

    add_h2(doc, "Sample AVIF picture element")
    add_code_block(doc,
        '<picture>\n'
        '  <source type="image/avif" srcset="/assets/hero.avif">\n'
        '  <source type="image/webp" srcset="/assets/hero.webp">\n'
        '  <img src="/assets/hero.jpg" alt="Stainless steel motor tubes manufactured at Apurvi\'s Mehsana plant"\n'
        '       width="1600" height="900" fetchpriority="high">\n'
        '</picture>'
    )

    add_page_break(doc)

    # 8. INDEXING / SEARCH CONSOLE / BING WMT
    add_h1(doc, "8. Indexing — GSC + Bing WMT + IndexNow")

    add_h2(doc, "Google Search Console")
    add_bullets(doc, [
        ("Verify Domain property", "Not URL prefix — apurviind.com as Domain. With hreflang in place, GSC auto-segments performance by country."),
        ("Submit sitemap.xml", "Plus /us/sitemap.xml and /uk/sitemap.xml when those folders exist."),
        ("URL Inspection → Request Indexing", "For every new page on Day 1."),
        ("Performance report", "Filter by country US, then UK. Track impressions weekly."),
    ])

    add_h2(doc, "Bing Webmaster Tools (CRITICAL for 2026)")
    add_callout(
        doc,
        "Why Bing matters in 2026",
        "Bing has ~12% US desktop share AND powers ChatGPT Search + Bing Copilot. The new AI Performance report (released Feb 11, 2026) is the only first-party way to see "
        "how often pages get cited in Microsoft Copilot. B2B buyers inside Microsoft 365 increasingly query Copilot for procurement — highest-ROI engine in 2026 for B2B.",
    )
    add_bullets(doc, [
        "Verify domain in Bing Webmaster Tools",
        "Submit sitemap.xml (+ /us/ + /uk/ sitemaps later)",
        "Check AI Performance weekly — track Copilot citations",
        "Configure crawl settings (allow site crawl, exclude /private/, /admin.php, /send-inquiry.php)",
    ])

    add_h2(doc, "IndexNow (already configured)")
    add_para(doc, "IndexNow key file is already at https://apurviind.com/9c0159795a04f63aee2e0d1a2be2401d.txt. To ping Bing + Yandex + Naver + Seznam + Yep on every publish:")
    add_code_block(doc,
        'curl "https://api.indexnow.org/indexnow?url=https://apurviind.com/your-new-page&key=9c0159795a04f63aee2e0d1a2be2401d"\n'
        '\n'
        '# Or batch POST:\n'
        'curl -X POST -H "Content-Type: application/json" \\\n'
        '  -d \'{"host":"apurviind.com","key":"9c0159795a04f63aee2e0d1a2be2401d",\n'
        '       "urlList":["https://apurviind.com/us/","https://apurviind.com/uk/"]}\' \\\n'
        '  "https://api.indexnow.org/indexnow"'
    )
    add_callout(
        doc,
        "Wire this into the build/publish pipeline",
        "Every time a page is added or content materially changes, fire an IndexNow ping. 22% of Bing-clicked URLs in Feb 2026 came via IndexNow. Google does NOT support IndexNow — use GSC URL Inspection for Google.",
    )

    add_page_break(doc)

    # 9. LEAD MAGNET ENGINEERING
    add_h1(doc, "9. Lead Magnet Technical Implementation")

    add_h2(doc, "Lead Magnet #1 — Pump-Tube Sizing Calculator (HIGHEST PRIORITY)")
    add_bullets(doc, [
        ("Page", "/tools/pump-tube-sizing-calculator/"),
        ("Inputs", "Pump HP (1–200), well depth (m), borehole ID (in), water type (fresh/brackish/saline)"),
        ("Outputs", "Recommended tube OD/ID, wall thickness, length, grade (304L/316L), MTC requirement"),
        ("Tech", "Vanilla JS or Alpine.js — no framework needed. Static HTML + JSON config file with sizing rules."),
        ("Embed widget", "Generate an embeddable <iframe src='https://apurviind.com/tools/pump-tube-sizing-calculator/embed?theme=light'> for 3rd-party pump distributor sites. Each embed = backlink."),
        ("Schema", "WebApplication + HowTo schema. Make calculator state shareable via URL params (improves social shares)."),
    ])

    add_h2(doc, "Lead Magnet #2 — EN 10204 3.1 MTC Annotated Sample (Day 6)")
    add_bullets(doc, [
        ("Page", "/downloads/en-10204-3-1-mtc-sample/"),
        ("File", "PDF (anonymised real Apurvi MTC with arrows/callouts explaining each field)"),
        ("Gate", "Email gate via send-inquiry.php (or Mailchimp/HubSpot form)"),
        ("Schema", "DigitalDocument schema, file accessible to AI crawlers (no robots.txt block)"),
    ])

    add_h2(doc, "Lead Magnet #4 — Vapour Recovery Compliance Map")
    add_bullets(doc, [
        ("Page", "/tools/vapor-recovery-compliance-map/"),
        ("Tech", "Interactive SVG map of US states (D3.js or vanilla) + UK regions. Click a state → show EPA Stage I/II + material requirements."),
        ("Data source", "Static JSON file with state-by-state EPA regulations + UK PPC regional rules"),
        ("Schema", "Dataset schema + WebApplication"),
    ])

    add_h2(doc, "Lead Magnets #5–#10")
    add_bullets(doc, [
        "#5 — Indian SS Pipe Landed-Cost Calculator: Vanilla JS, Section 232 + CBAM rules baked in",
        "#6 — Submersible Motor Tube Spec Sheet Bundle: 8 PDFs zipped, email-gated",
        "#7 — 12-Point Supplier Audit Checklist: .docx download",
        "#8 — Wall Thickness / Collapse Pressure Workbook: .xlsx with pre-loaded formulas",
        "#9 — Pitting Corrosion Field Study: 12-page PDF whitepaper (primary-source research = high AI citation potential)",
        "#10 — Standards Cross-Reference Card: Printable A3 PDF",
    ])

    add_page_break(doc)

    # 10. NEW URLS TO BUILD
    add_h1(doc, "10. New URLs to Build")
    add_para(doc, "Below is the full list of new pages to create — owners should be assigned per URL. Total ~50 new pages.")

    add_h2(doc, "US-targeted landing pages")
    add_bullets(doc, [
        "/us/",
        "/us/products",
        "/us/industries",
        "/us/contact",
        "/us/company",
        "/us/astm-a312-tp304-stainless-steel-pipe-supplier",
        "/us/astm-a268-stainless-steel-tubing-manufacturer",
        "/us/astm-a213-boiler-superheater-tubing-india",
        "/us/submersible-pump-motor-shell-tube-supplier",
        "/us/aluminum-vapor-recovery-tubing-stage-i-retrofit",
        "/us/non-china-stainless-steel-tube-source-2026",
        "/us/section-232-tariff-impact-indian-tubing-buyer-guide",
        "/us/high-hp-submersible-motor-tube",
        "/us/stainless-steel-tube-for-oem",
    ])

    add_h2(doc, "UK-targeted landing pages")
    add_bullets(doc, [
        "/uk/",
        "/uk/products",
        "/uk/industries",
        "/uk/contact",
        "/uk/company",
        "/uk/en-10296-2-welded-stainless-steel-tube-supplier",
        "/uk/en-10204-3-1-mtc-stainless-pipe-manufacturer",
        "/uk/ss-304-316-seamless-tube-uk-distributor",
        "/uk/pump-motor-tubing-uk-water-industry",
        "/uk/aluminium-vapour-recovery-tubing",
        "/uk/precision-aluminium-tubing",
        "/uk/borehole-pump-tube-supplier",
    ])

    add_h2(doc, "Comparison + content marketing pages")
    add_bullets(doc, [
        "/comparison/india-vs-china-ss-pipe-sourcing-2026",
        "/comparison/astm-a268-vs-a269-vs-a312-which-spec",
        "/comparison/en-10296-vs-astm-a554-tolerance",
        "/guides/en-10204-3-1-mtc-walkthrough-pdf",
        "/guides/how-us-buyers-qualify-indian-tube-manufacturer",
        "/guides/submersible-motor-shell-material-selection-304-vs-316",
    ])

    add_h2(doc, "Industry vertical landing pages (mirror Plymouth Tube's playbook)")
    add_bullets(doc, [
        "/industries/water-well-submersible-pump-oem",
        "/industries/petroleum-fueling-vapor-recovery",
        "/industries/oil-gas-process-piping",
        "/industries/chemical-processing-en-compliant",
    ])

    add_h2(doc, "Tools / lead magnet pages")
    add_bullets(doc, [
        "/tools/pump-tube-sizing-calculator/",
        "/tools/vapor-recovery-compliance-map/",
        "/tools/landed-cost-calculator/",
        "/tools/astm-comparison/",
        "/downloads/en-10204-3-1-mtc-sample/",
        "/downloads/supplier-audit-checklist/",
        "/downloads/wall-thickness-workbook/",
        "/downloads/pitting-corrosion-whitepaper/",
        "/downloads/standards-cross-reference/",
        "/downloads/motor-tube-spec-bundle/",
    ])

    add_page_break(doc)

    # 11. EVERY PAGE CHECKLIST
    add_h1(doc, "11. Every New Page — Required Elements Checklist")
    add_table(
        doc,
        ["Element", "US (/us/*)", "UK (/uk/*)", "Comparison/Guides", "Industry"],
        [
            ["<html lang='..'> attr", "en-US", "en-GB", "en", "en"],
            ["Canonical (self-reference)", "Required", "Required", "Required", "Required"],
            ["Hreflang block (all 4 alternates)", "Required", "Required", "Required (canonical only)", "Required (canonical only)"],
            ["geo.* meta tags", "REMOVE", "REMOVE", "REMOVE", "REMOVE"],
            ["Title tag (no 'India'/'Ahmedabad')", "US-buyer focused", "UK-buyer focused", "Neutral", "Industry-focused"],
            ["Meta description with $/£ + port name", "Yes", "Yes", "Optional", "Optional"],
            ["Schema: Organization + Service + Breadcrumb", "Yes", "Yes", "Yes (no Service)", "Yes"],
            ["Schema: Article/FAQPage", "If blog", "If blog", "Yes", "Yes"],
            ["Open Graph + Twitter Card", "Yes", "Yes", "Yes", "Yes"],
            ["Author byline + Person schema", "If blog", "If blog", "Yes", "No"],
            ["dateModified updated on every change", "Yes", "Yes", "Yes", "Yes"],
            ["EXIF-tagged original photos", "Yes", "Yes", "Yes", "Yes"],
            ["Primary source citations (ASTM/BSI/EPA)", "If technical", "If technical", "Yes", "Yes"],
            ["Calendly + US/UK phone in footer", "Yes", "Yes", "Yes", "Yes"],
            ["RFQ form on page", "Yes", "Yes", "Yes", "Yes"],
            ["Self-host fonts + AVIF hero image", "Yes", "Yes", "Yes", "Yes"],
        ],
    )

    add_page_break(doc)

    # 12. SPRINT PLAN
    add_h1(doc, "12. Tech Sprint Plan (30 Days)")

    add_h2(doc, "Sprint 1 (Days 1–7) — Foundation")
    add_table(
        doc,
        ["Day", "Task", "Owner"],
        [
            ["1", "Create /us/ folder with index.html cloned + US-locale hreflang + schema", "Frontend dev"],
            ["1", "Create /uk/ folder with index.html cloned + UK-locale hreflang + schema", "Frontend dev"],
            ["2", "Remove geo.* meta tags from /us/* and /uk/* pages", "Frontend dev"],
            ["2", "Update /sitemap.xml with /us/ and /uk/ URLs + xhtml:link hreflang entries", "Frontend dev"],
            ["3", "Wire IndexNow ping into deploy hook (Vercel/Netlify/cron)", "Backend dev"],
            ["3", "Clone /us/products + /uk/products with US/UK spelling localised", "Content + frontend"],
            ["4", "Clone /us/industries + /uk/industries", "Content + frontend"],
            ["4", "Clone /us/contact + /uk/contact with Calendly + US/UK phone placeholders", "Frontend dev"],
            ["5", "Self-host fonts (subset Latin) + add preload tags", "Frontend dev"],
            ["5", "Convert hero images to AVIF + WebP fallback", "Frontend dev"],
            ["6", "Add Person schema + author byline to existing 3 blog posts", "Frontend dev"],
            ["7", "Refactor existing blog posts to AEO template structure", "Content + frontend"],
        ],
    )

    add_h2(doc, "Sprint 2 (Days 8–14) — Tools / Lead Magnets")
    add_table(
        doc,
        ["Day", "Task", "Owner"],
        [
            ["8",  "Build /tools/pump-tube-sizing-calculator/ (vanilla JS + JSON config + iframe-embed mode)", "Frontend dev"],
            ["10", "Build /tools/vapor-recovery-compliance-map/ (interactive SVG D3.js + static JSON data)", "Frontend dev"],
            ["11", "Build /tools/landed-cost-calculator/ (JS, S232 + CBAM rules baked in)", "Frontend dev"],
            ["12", "Build /downloads/* PDF gating pages (form → email → download link)", "Backend dev"],
            ["13", "Add WebApplication + HowTo schema to all /tools/ pages", "Frontend dev"],
            ["14", "Add DigitalDocument schema to all /downloads/ pages", "Frontend dev"],
        ],
    )

    add_h2(doc, "Sprint 3 (Days 15–21) — Vertical Pages + Comparisons")
    add_table(
        doc,
        ["Day", "Task", "Owner"],
        [
            ["15-17", "Build 4 /industries/* vertical landing pages", "Content + frontend"],
            ["18-19", "Build 3 /comparison/* pages (India vs China, ASTM A268 vs A312, EN vs ASTM)", "Content + frontend"],
            ["20-21", "Build 3 /guides/* deep guides (MTC walkthrough, qualify Indian mfg, motor shell selection)", "Content + frontend"],
        ],
    )

    add_h2(doc, "Sprint 4 (Days 22–30) — Long-Tail Sweep + Polish")
    add_table(
        doc,
        ["Day", "Task", "Owner"],
        [
            ["22-24", "Build remaining /us/* + /uk/* product-specific landing pages (~12 pages)", "Content + frontend"],
            ["25",    "Lighthouse audit all new pages — push every score >90", "Frontend dev"],
            ["26",    "Run W3C HTML validator, schema.org validator, rich results test on every new URL", "Frontend dev"],
            ["27",    "Verify hreflang bidirectional balance with Ahrefs / Sitebulb / Screaming Frog audit", "SEO dev"],
            ["28",    "Wire AI citation tracking — Otterly or Profound API for AI Overview monitoring", "Backend dev"],
            ["29",    "Final crawl + fix any 404s, broken internal links, missing canonicals", "Frontend dev"],
            ["30",    "Generate technical SEO report — Lighthouse scores, schema validations, hreflang audit", "SEO dev"],
        ],
    )

    add_page_break(doc)

    # 13. TESTING CHECKLIST
    add_h1(doc, "13. Testing Checklist Before Each Page Ships")
    add_numbered(doc, [
        "HTML validates at validator.w3.org (zero errors)",
        "Schema validates at validator.schema.org + Google Rich Results test",
        "Lighthouse scores ≥ 90 on mobile (Performance, Accessibility, SEO, Best Practices)",
        "INP measured via DevTools Performance panel < 200ms",
        "Hreflang bidirectional balance — every alternate references back",
        "Canonical self-references (does NOT point to a different region)",
        "Open Graph image renders correctly (test in opengraph.xyz)",
        "Twitter Card preview renders (cards-dev.twitter.com)",
        "Page loads in < 2.5s on Slow 3G throttling (WebPageTest.org)",
        "Mobile tap targets verified ≥ 48x48px (DevTools mobile emulation)",
        "Colour contrast checker passes 4.5:1 (Lighthouse a11y)",
        "No console errors in DevTools",
        "Sitemap.xml updated and submitted to GSC + Bing WMT",
        "IndexNow ping fired",
        "Internal links from existing 16 pages added (crawl-depth ≤ 3)",
    ])

    add_page_break(doc)

    # 14. COMMUNICATION
    add_h1(doc, "14. Cross-Team Communication Points (Tech ↔ SEO)")
    add_callout(
        doc,
        "Hand-off points where the SEO team depends on tech",
        "(1) Hreflang shipped before SEO team starts paid campaigns. "
        "(2) /us/ /uk/ folders LIVE before any directory listings reference these URLs. "
        "(3) Lead magnet pages LIVE before SEO team starts content marketing. "
        "(4) IndexNow + Bing WMT setup BEFORE blog publish cadence starts.",
    )
    add_callout(
        doc,
        "What the tech team needs from SEO",
        "(1) Final approved page titles per URL (US-buyer + UK-buyer language). "
        "(2) Author bios + LinkedIn URLs for Person schema. "
        "(3) US 1-800 + UK 0800 numbers to wire into footer/contact pages. "
        "(4) Approved hero images for /us/ /uk/ — must be original Apurvi photos with rights cleared.",
    )

    # 15. APPENDIX
    add_h1(doc, "15. Appendix — Quick-Reference Snippets")

    add_h2(doc, "<head> snippet bundle for any /us/* page")
    add_code_block(doc,
        '<link rel="canonical" href="https://apurviind.com/us/" />\n'
        '<link rel="alternate" hreflang="en-us" href="https://apurviind.com/us/" />\n'
        '<link rel="alternate" hreflang="en-gb" href="https://apurviind.com/uk/" />\n'
        '<link rel="alternate" hreflang="en-in" href="https://apurviind.com/" />\n'
        '<link rel="alternate" hreflang="x-default" href="https://apurviind.com/" />\n'
        '\n'
        '<meta name="description" content="US buyers: Apurvi Industries supplies ASTM A312 / A268 stainless steel motor tubes, SS 304 / SS 316 pressure-pump sleeves, and aluminum vapor recovery tubing to pump OEMs, water-treatment integrators, and fueling-equipment manufacturers. EN 10204 3.1 mill certs. Container-load export from Mumbai/Mundra to Houston, LA, NY/NJ. Reply in one working day." />\n'
        '<meta name="robots" content="index, follow, max-snippet:-1, max-image-preview:large, max-video-preview:-1" />\n'
        '\n'
        '<!-- LCP image hint -->\n'
        '<link rel="preload" as="image" href="/us/assets/hero-ss-motor-tubes.avif" fetchpriority="high" />\n'
        '\n'
        '<!-- Self-hosted font -->\n'
        '<link rel="preload" href="/assets/fonts/inter-var.woff2" as="font" type="font/woff2" crossorigin />'
    )

    add_h2(doc, "Sitemap-level hreflang block for sitemap.xml")
    add_code_block(doc,
        '<url>\n'
        '  <loc>https://apurviind.com/us/</loc>\n'
        '  <lastmod>2026-06-29</lastmod>\n'
        '  <xhtml:link rel="alternate" hreflang="en-us" href="https://apurviind.com/us/" />\n'
        '  <xhtml:link rel="alternate" hreflang="en-gb" href="https://apurviind.com/uk/" />\n'
        '  <xhtml:link rel="alternate" hreflang="en-in" href="https://apurviind.com/" />\n'
        '  <xhtml:link rel="alternate" hreflang="x-default" href="https://apurviind.com/" />\n'
        '</url>'
    )

    doc.save("/Users/sayujpillai/Desktop/apurvi/Apurviind.com/Apurvi-Tech-Team-Brief.docx")
    print("Tech doc saved.")


if __name__ == "__main__":
    build_seo_doc()
    build_tech_doc()
    print("Both docs generated successfully.")
